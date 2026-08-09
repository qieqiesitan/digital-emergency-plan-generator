"""导入文件解析：xlsx / csv / docx / pdf / txt → 纯文本。"""
import csv
import io


def parse_file_text(filename: str, data: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "csv":
            return _parse_csv(data)
        if ext == "xlsx":
            return _parse_xlsx(data)
        if ext == "docx":
            return _parse_docx(data)
        if ext == "pdf":
            return _parse_pdf(data)
        if ext in ("txt", "md"):
            return data.decode("utf-8", errors="ignore")
    except Exception as e:
        raise ValueError("文件解析失败，请确认文件未损坏且格式正确") from e
    raise ValueError(f"不支持的文件格式：.{ext}，支持 xlsx/csv/docx/pdf/txt")


def _parse_csv(data: bytes) -> str:
    raw = None
    for enc in ("utf-8-sig", "gbk"):
        try:
            raw = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if raw is None:
        raw = data.decode("utf-8", errors="ignore")
    rows = list(csv.reader(io.StringIO(raw)))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows if any(cell.strip() for cell in row))


def _parse_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        parts = []
        for ws in wb.worksheets:
            parts.append(f"【工作表：{ws.title}】")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    finally:
        wb.close()


def _parse_docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()
