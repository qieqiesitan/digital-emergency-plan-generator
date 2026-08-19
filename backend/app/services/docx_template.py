"""
应急预案 DOCX 模板引擎
基于 紫楹台酒店应急预案.docx 参考格式，提供完整的 Word 文档生成能力。

格式体系（10种样式）：
  - 封面标题 (CoverTitle):   宋体 26pt 居中 — 封面单位名称/文档标题
  - 封面落款 (CoverSign):    黑体/仿宋 18pt 居中 — 封面落款/批准页
  - 正文大标题 (BodyTitle):   宋体 22pt 居中 — 正文页主标题
  - Normal:                 仿宋 16pt 两端对齐，首行缩进2字符
  - Heading 1:             黑体 16pt 加粗 — 一级标题
  - Heading 2:             宋体 16pt 加粗 — 二级标题
  - Heading 3:             宋体 16pt 加粗 — 三级标题
  - Heading 4:             仿宋 16pt 加粗 — 四级标题
  - Body Text:             仿宋_GB2312 — 特殊正文
  - IDX-B:                 華康中楷體 14pt — 索引专用
"""

import io
import re
import logging
import traceback
from datetime import datetime
from typing import Optional, Callable

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls

from docx.oxml import parse_xml
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════
# 单位换算常量
# ═══════════════════════════════════════════
EMU_PER_PT = 12700
EMU_PER_CM = 360000

# ═══════════════════════════════════════════
# 样式常量（与参考文档一致）
# ═══════════════════════════════════════════
STYLE_COVER_TITLE = "Cover Title"
STYLE_COVER_SIGN = "Cover Sign"
STYLE_BODY_TITLE = "Body Title"
STYLE_IDX_B = "IDX-B"

# 字体
FONT_SONGTI = "宋体"
FONT_HEITI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_FANGSONG_GB = "仿宋_GB2312"
FONT_KAITI = "楷体"
FONT_TNR = "Times New Roman"
FONT_HK_ZHONGKAI = "華康中楷體"

# 字号 (Pt)
SIZE_COVER_TITLE = 26       # 封面标题
SIZE_COVER_SIGN = 18        # 封面落款
SIZE_BODY_TITLE = 22        # 正文大标题
SIZE_HEADING = 16           # 标题 1-4
SIZE_NORMAL = 16            # 正文
SIZE_IDX = 14               # 索引

# 首行缩进 (Pt) — 2字符 ≈ 32pt @ 16pt字号
FIRST_INDENT_NORMAL = 32    # Normal 首行缩进
FIRST_INDENT_HEADING = 31.6 # 标题首行缩进

# 页边距 (Cm)
MARGIN_COVER_LEFT = 2.8
MARGIN_COVER_RIGHT = 2.6
MARGIN_COVER_TOP = 3.7
MARGIN_COVER_BOTTOM = 3.5

MARGIN_BODY_LEFT = 3.18
MARGIN_BODY_RIGHT = 3.18
MARGIN_BODY_TOP = 2.54
MARGIN_BODY_BOTTOM = 2.54

# 编号格式定义
NUM_FMT_CHINESE = "chineseCountingThousand"     # 一、二、三
NUM_FMT_PAREN = "chineseCountingThousand"        # （一）（二）
NUM_FMT_ARABIC = "decimal"                       # 1. 2. 3.
NUM_FMT_ARABIC_PAREN = "decimal"                 # (1) (2)


# ═══════════════════════════════════════════
# 样式注册
# ═══════════════════════════════════════════

def _define_style(doc, style_name, base_style="Normal", font_name=None,
                  font_size=None, bold=None, alignment=None,
                  first_line_indent=None, color=None, space_before=None, space_after=None):
    """定义或获取样式。若样式已存在则返回，否则创建。"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
    if base_style:
        style.base_style = doc.styles[base_style]
    fmt = style.paragraph_format
    if alignment is not None:
        fmt.alignment = alignment
    if first_line_indent is not None:
        fmt.first_line_indent = Pt(first_line_indent)
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    font = style.font
    if font_name:
        font.name = font_name
    if font_size:
        font.size = Pt(font_size)
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = color
    return style


def register_all_styles(doc: Document):
    """注册所有自定义样式。调用一次即可。"""
    # ── Normal 基准 ──
    normal = doc.styles["Normal"]
    normal.font.name = FONT_TNR
    normal.font.size = Pt(SIZE_NORMAL)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    _set_full_fonts(normal, ascii_font=FONT_TNR, ea_font=FONT_FANGSONG)

    # ── 封面标题 ──
    _define_style(doc, STYLE_COVER_TITLE,
                  font_name=FONT_SONGTI, font_size=SIZE_COVER_TITLE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=0, space_after=0)
    _set_east_asian_font(doc.styles[STYLE_COVER_TITLE], FONT_SONGTI)

    # ── 封面落款 ──
    _define_style(doc, STYLE_COVER_SIGN,
                  font_name=FONT_FANGSONG, font_size=SIZE_COVER_SIGN,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=0, space_after=0)
    _set_east_asian_font(doc.styles[STYLE_COVER_SIGN], FONT_FANGSONG)

    # ── 正文大标题 ──
    _define_style(doc, STYLE_BODY_TITLE,
                  font_name=FONT_SONGTI, font_size=SIZE_BODY_TITLE,
                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=Pt(24), space_after=Pt(12))
    _set_east_asian_font(doc.styles[STYLE_BODY_TITLE], FONT_SONGTI)

    # ── Heading 1 ──
    h1 = _define_style(doc, "Heading 1",
                       font_name=FONT_HEITI, font_size=SIZE_HEADING,
                       bold=True, first_line_indent=FIRST_INDENT_HEADING,
                       space_before=Pt(28), space_after=Pt(0))
    _set_east_asian_font(h1, FONT_HEITI)

    # ── Heading 2 ──
    h2 = _define_style(doc, "Heading 2",
                       font_name=FONT_SONGTI, font_size=SIZE_HEADING,
                       bold=True, first_line_indent=FIRST_INDENT_HEADING,
                       space_before=Pt(6), space_after=Pt(0))
    _set_east_asian_font(h2, FONT_SONGTI)

    # ── Heading 3 ──
    h3 = _define_style(doc, "Heading 3",
                       font_name=FONT_SONGTI, font_size=SIZE_HEADING,
                       bold=True, first_line_indent=FIRST_INDENT_HEADING,
                       space_before=Pt(6), space_after=Pt(0))
    _set_east_asian_font(h3, FONT_SONGTI)

    # ── Heading 4 ──
    h4 = _define_style(doc, "Heading 4",
                       font_name=FONT_FANGSONG, font_size=SIZE_HEADING,
                       bold=True, first_line_indent=FIRST_INDENT_HEADING,
                       space_before=Pt(6), space_after=Pt(0))
    _set_east_asian_font(h4, FONT_FANGSONG)


    # ── Heading 5 ──
    h5 = _define_style(doc, "Heading 5",
                       font_name=FONT_FANGSONG, font_size=SIZE_HEADING,
                       bold=True, first_line_indent=FIRST_INDENT_HEADING,
                       space_before=Pt(6), space_after=Pt(0),
                       color=RGBColor(0, 0, 0))
    _set_east_asian_font(h5, FONT_FANGSONG)

    # ── Heading 6 ──
    h6 = _define_style(doc, "Heading 6",
                       font_name=FONT_FANGSONG, font_size=SIZE_HEADING,
                       bold=True, first_line_indent=FIRST_INDENT_HEADING,
                       space_before=Pt(6), space_after=Pt(0),
                       color=RGBColor(0, 0, 0))
    _set_east_asian_font(h6, FONT_FANGSONG)    # ── Body Text ──
    _define_style(doc, "Body Text",
                  font_name=FONT_FANGSONG_GB, font_size=SIZE_NORMAL,
                  first_line_indent=FIRST_INDENT_NORMAL)
    _set_east_asian_font(doc.styles["Body Text"], FONT_FANGSONG_GB)

    # ── IDX-B ──
    _define_style(doc, STYLE_IDX_B,
                  font_name=FONT_HK_ZHONGKAI, font_size=SIZE_IDX,
                  first_line_indent=0)
    _set_east_asian_font(doc.styles[STYLE_IDX_B], FONT_HK_ZHONGKAI)

    logger.info("All custom styles registered")


def _set_east_asian_font(style, font_name):
    """为样式设置东亚字体（通过 XML 操作）。"""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


# ═══════════════════════════════════════════
# 页面设置工具
# ═══════════════════════════════════════════

def set_page_margins(section, left_cm, right_cm, top_cm, bottom_cm):
    """设置页面边距（厘米）。"""
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(right_cm)
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)


def add_section(doc: Document, left_cm=None, right_cm=None,
                top_cm=None, bottom_cm=None):
    """添加新节（Section），可选自定义页边距。"""
    new_section = doc.add_section()
    if left_cm:
        set_page_margins(new_section, left_cm, right_cm or left_cm,
                         top_cm or 2.54, bottom_cm or 2.54)
    return new_section


# ═══════════════════════════════════════════
# 封面构建
# ═══════════════════════════════════════════

def build_cover(doc: Document, *,
                plan_number: str = "",
                version_number: str = "",
                company_name: str,
                doc_title: str = "生产安全事故应急预案",
                signature_company: str = "",
                signature_date: str = "",
                approval_texts: list[str] | None = None):
    """构建封面页和批准页。

    封面页结构：
      1. 编号/版本号行（右上角对齐）
      2. 公司名称（封面标题）
      3. 文档标题（封面标题）
      4. 落款公司名 + 日期（封面落款）
      5. 批准页文字
    """
    # ── 封面页 ──
    # 编号/版本号
    if plan_number or version_number:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"应急预案编号：{plan_number}　　　　应急预案版本号：{version_number}")
        r.font.name = FONT_SONGTI
        r.font.size = Pt(14)
        r.bold = True
        _set_east_asian_font_in_run(r, FONT_SONGTI)

    # 空行撑开
    for _ in range(3):  # reduced from 6
        doc.add_paragraph("")

    # 公司名称
    p = doc.add_paragraph(company_name, style=STYLE_COVER_TITLE)

    # 文档标题
    p = doc.add_paragraph(doc_title, style=STYLE_COVER_TITLE)

    # 空行
    for _ in range(2):  # reduced from 5
        doc.add_paragraph("")

    # 落款
    sig_company = signature_company or company_name
    p = doc.add_paragraph(sig_company, style=STYLE_COVER_SIGN)
    for run in p.runs:
        run.font.name = FONT_HEITI
        _set_east_asian_font_in_run(run, FONT_HEITI)

    sig_date = signature_date or datetime.now().strftime("%Y年%m月")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sig_date)
    r.font.name = FONT_HEITI
    r.font.size = Pt(SIZE_COVER_SIGN)
    _set_east_asian_font_in_run(r, FONT_HEITI)

    # ── 批准页（紧随封面后） ──
    doc.add_page_break()

    p = doc.add_paragraph("批准页", style=STYLE_COVER_SIGN)
    for run in p.runs:
        run.font.name = FONT_SONGTI
        run.bold = True
        _set_east_asian_font_in_run(run, FONT_SONGTI)
    # 设置间距
    p.paragraph_format.space_after = Pt(24)

    if approval_texts:
        for text in approval_texts:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Pt(FIRST_INDENT_NORMAL)
            _set_run_font(p, FONT_FANGSONG, SIZE_NORMAL)
    else:
        default_texts = [
            f"为贯彻落实《中华人民共和国安全生产法》及其他法律法规和有关文件要求，结合{company_name}的安全生产实际和危险性分析及本行业事故案例情况内容，依据《中华人民共和国突发事件应对法》（中华人民共和国主席令第69号）《生产经营单位生产安全事故应急预案编制导则》（GB/T 29639-2020）《生产安全事故应急预案管理办法》（国家应急管理部令第2号）《生产安全事故应急条例》（国务院令第708号）",
            f"《生产安全事故应急预案》由《综合应急预案》《专项应急预案》和《现场处置方案》组成，我公司邀请相关专家召开了专家评审会议，并通过专家评审，现予以发布，自发布之日起实施。",
            "望各部门及全体员工认真学习该预案内容，熟悉预案程序，并严格遵守执行。",
        ]
        for text in default_texts:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Pt(FIRST_INDENT_NORMAL)
            _set_run_font(p, FONT_FANGSONG, SIZE_NORMAL)

    # 签字区
    doc.add_paragraph("")
    p = doc.add_paragraph("主要负责人（签字）：")
    p.paragraph_format.first_line_indent = Pt(31)
    _set_run_font(p, FONT_FANGSONG, SIZE_NORMAL)

    p = doc.add_paragraph("日期：　　　年　　　月　　　日")
    p.paragraph_format.first_line_indent = Pt(31)
    _set_run_font(p, FONT_FANGSONG, SIZE_NORMAL)


def _set_run_font(paragraph, font_name, font_size_pt, bold=False, color=None):
    """设置段落首个 run 的字体。若无 run 则创建。"""
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run(paragraph.text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    _set_east_asian_font_in_run(run, font_name)

def _set_full_fonts(style, ascii_font, ea_font):
    """设置样式的完整字体属性（拉丁 + 东亚）。"""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), ea_font)
    rfonts.set(qn("w:cs"), ea_font)

def _set_east_asian_font_in_run(run, font_name):
    """为单个 run 设置东亚字体。"""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


# ═══════════════════════════════════════════
# 签署页
# ═══════════════════════════════════════════

def build_signature_page(doc: Document, signers: list[dict]):
    """构建应急预案执行部门签署页。

    signers: [{"seq": 1, "name": "张红", "title": "总经理"}, ...]
    """
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("应急预案执行部门签署页")
    r.font.name = FONT_SONGTI
    r.font.size = Pt(18)
    r.bold = True
    _set_east_asian_font_in_run(r, FONT_SONGTI)
    p.paragraph_format.space_after = Pt(24)

    if not signers:
        return

    table = doc.add_table(rows=len(signers) + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["序号", "姓名", "职务", "签署"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = FONT_SONGTI
                r.font.size = Pt(12)
                r.bold = True
                _set_east_asian_font_in_run(r, FONT_SONGTI)

    # 数据行
    for i, signer in enumerate(signers):
        for j, key in enumerate(["seq", "name", "title", ""]):
            cell = table.cell(i + 1, j)
            cell.text = str(signer.get(key, "")) if key else ""
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = FONT_FANGSONG
                    r.font.size = Pt(12)
                    _set_east_asian_font_in_run(r, FONT_FANGSONG)


# ═══════════════════════════════════════════
# 正文内容构建
# ═══════════════════════════════════════════

def add_body_title(doc: Document, title: str):
    """添加正文大标题（如"综合应急预案"）。"""
    p = doc.add_paragraph(title, style=STYLE_BODY_TITLE)


def add_heading(doc: Document, text: str, level: int):
    """添加标题，自动映射 level 到 Heading 1-6。"""
    h = doc.add_heading(text, level=min(level, 6))
    # 确保首行缩进
    if h.paragraph_format.first_line_indent is None:
        h.paragraph_format.first_line_indent = Pt(FIRST_INDENT_HEADING)
    return h


def add_normal_paragraph(doc: Document, text: str):
    """添加首行缩进的正文段落。"""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(FIRST_INDENT_NORMAL)
    return p


def add_numbered_paragraph(doc: Document, text: str, level: int = 0):
    """添加带编号的段落。"""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(FIRST_INDENT_NORMAL)
    # 编号由 AI 生成的内容自带（如 1. 2. 或 ① ②）
    return p


# ═══════════════════════════════════════════
# 表格构建
# ═══════════════════════════════════════════

def build_table(doc: Document, headers: list[str], rows: list[list[str]],
                col_widths: list[float] | None = None):
    """构建标准格式表格。

    col_widths: 列宽列表（厘米），为空则自动均分。
    """
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = FONT_SONGTI
                r.font.size = Pt(10.5)
                r.bold = True
                _set_east_asian_font_in_run(r, FONT_SONGTI)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j >= len(headers):
                break
            cell = table.cell(i + 1, j)
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = FONT_FANGSONG
                    r.font.size = Pt(10.5)
                    _set_east_asian_font_in_run(r, FONT_FANGSONG)

    # 列宽
    if col_widths:
        for j, width in enumerate(col_widths):
            if j < len(headers):
                for row in table.rows:
                    row.cells[j].width = Cm(width)

    return table


# ═══════════════════════════════════════════
# HTML/Markdown → DOCX 转换（增强版）
# ═══════════════════════════════════════════

def html_to_docx_content(doc: Document, html_content: str, base_level: int = 1, _depth: int = 0):
    # ponytail: prevent infinite recursion from nested divs
    if _depth > 10:
        p = doc.add_paragraph()
        p.add_run(html_content[:200] + "..." if len(html_content) > 200 else html_content)
        return
    """将 HTML 内容转换为 docx 段落，保留富文本格式。

    支持：h1-h6, p, table, ul/ol, blockquote, hr, pre, div
    base_level: 起始 heading 级别偏移
    """
    if not html_content or not html_content.strip():
        return

    # 如果内容是纯文本（非 HTML），先转 Markdown 再按 HTML 处理
    if not html_content.strip().startswith("<"):
        import markdown as _md2
        html_content = _md2.markdown(html_content, extensions=["tables", "fenced_code", "md_in_html"])

    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.children:
        if isinstance(element, str):
            text = str(element).strip()
            if text:
                add_normal_paragraph(doc, text)
            continue

        tag = element.name
        if tag is None:
            continue

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1]) + base_level - 1
            level = min(max(level, 1), 6)
            add_heading(doc, element.get_text().strip(), level)

        elif tag == "p":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(FIRST_INDENT_NORMAL)
            _add_inline_runs(p, element)

        elif tag == "table":
            rows = element.find_all("tr")
            if not rows:
                continue
            # 确定列数
            max_cols = 0
            for row in rows:
                cols = len(row.find_all(["th", "td"]))
                max_cols = max(max_cols, cols)
            if max_cols == 0:
                continue

            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for j, cell in enumerate(cells):
                    if j >= max_cols:
                        break
                    doc_cell = table.cell(i, j)
                    doc_cell.text = cell.get_text().strip()
                    is_header = cell.name == "th" or i == 0
                    for cp in doc_cell.paragraphs:
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for cr in cp.runs:
                            cr.font.name = FONT_SONGTI if is_header else FONT_FANGSONG
                            cr.font.size = Pt(10.5)
                            if is_header:
                                cr.bold = True
                            _set_east_asian_font_in_run(cr, FONT_SONGTI if is_header else FONT_FANGSONG)

            doc.add_paragraph("")

        elif tag in ("ul", "ol"):
            for li in element.find_all("li", recursive=False):
                prefix = ""  # use Word List Bullet style, no manual bullet
                p = doc.add_paragraph(li.get_text().strip(), style="List Bullet")
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(-0.5)

        elif tag == "blockquote":
            p = doc.add_paragraph()
            run = p.add_run(element.get_text().strip())
            run.italic = True
            # run.font.color.rgb removed — all text black

        elif tag == "hr":
            doc.add_paragraph("-" * 60)

        elif tag == "pre":
            code_el = element.find("code")
            if code_el and "language-mermaid" in (code_el.get("class") or []):
                p = doc.add_paragraph()
                r = p.add_run("[Mermaid 流程图]")
                # r.font.color.rgb removed — all text black
                r.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                r = p.add_run(element.get_text())
                r.font.name = "Courier New"
                r.font.size = Pt(9)

        elif tag == "div":
            # 递归处理
            html_to_docx_content(doc, str(element), base_level, _depth + 1)

        else:
            text = element.get_text().strip()
            if text:
                add_normal_paragraph(doc, text)


def _add_inline_runs(paragraph, element):
    """将 HTML 元素内的内联格式（bold/italic/underline 等）写入段落。"""
    for child in element.children:
        if isinstance(child, str) or not hasattr(child, 'name'):
            text = str(child).strip('\n')
            if text:
                paragraph.add_run(text)
        elif child.name in ('strong', 'b'):
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ('em', 'i'):
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'u':
            run = paragraph.add_run(child.get_text())
            run.underline = True
        elif child.name in ('s', 'del'):
            run = paragraph.add_run(child.get_text())
            run.font.strike = True
        elif child.name == 'a':
            run = paragraph.add_run(child.get_text())
            # run.font.color.rgb removed — all text black
        elif child.name == 'br':
            paragraph.add_run('\n')
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
        elif child.name == 'span':
            style_attr = child.get('style', '')
            run = paragraph.add_run(child.get_text())
            if 'font-weight: bold' in style_attr or 'font-weight:700' in style_attr:
                run.bold = True
            if 'font-style: italic' in style_attr:
                run.italic = True
            if 'text-decoration: underline' in style_attr:
                run.underline = True
            if 'text-decoration: line-through' in style_attr:
                run.font.strike = True
            if 'background-color: yellow' in style_attr or 'background: yellow' in style_attr:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            _add_inline_runs(paragraph, child)


# ═══════════════════════════════════════════
# Markdown 表格预处理
# ═══════════════════════════════════════════

def fix_markdown_tables(html: str) -> str:
    """将 HTML <p> 标签内的 Markdown 管道表格转换为 <table> 元素。"""
    def convert(m):
        inner = m.group(1)
        lines = inner.strip().split('\n')
        clean = [l.strip() for l in lines if l.strip()]
        if len(clean) < 2:
            return m.group(0)

        sep_idx = -1
        for i, line in enumerate(clean):
            if re.match(r'^\|[\s\-:|]+\|', line):
                sep_idx = i
                break
        if sep_idx < 1:
            return m.group(0)

        header_line = clean[sep_idx - 1]
        leading = ''
        if not header_line.startswith('|'):
            pipe_pos = header_line.find('|')
            if pipe_pos > 0:
                leading = header_line[:pipe_pos].strip()
                header_line = header_line[pipe_pos:]

        headers = [c.strip() for c in header_line.split('|')[1:-1]]
        data_rows = []
        for line in clean[sep_idx + 1:]:
            if not line.startswith('|'):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                data_rows.append(cells)
        if not data_rows:
            return m.group(0)

        th = ''.join(f'<th>{h}</th>' for h in headers)
        td_rows = ''.join(
            '<tr>' + ''.join(f'<td>{c}</td>' for c in row) + '</tr>'
            for row in data_rows
        )
        table_html = f'<table><thead><tr>{th}</tr></thead><tbody>{td_rows}</tbody></table>'
        return f'<p>{leading}</p>{table_html}' if leading else table_html

    return re.sub(
        r'<p[^>]*>(.*?\|[\s\-:|]+\|.*?)</p>',
        convert,
        html,
        flags=re.DOTALL
    )


# ═══════════════════════════════════════════
# 主入口：生成完整预案文档
# ═══════════════════════════════════════════

def _build_docx_section_numbers(sections: list) -> dict:
    """为章节生成编号。level 0 → 1,2,3; level 1 → 1.1,1.2; level 2 → 1.1.1"""
    counters = [0] * 6
    numbers = {}
    for i, sec in enumerate(sections):
        level = sec.get("level", 0)
        counters[level] += 1
        for j in range(level + 1, len(counters)):
            counters[j] = 0
        parts = [str(counters[j]) for j in range(level + 1) if counters[j] > 0]
        numbers[i] = ".".join(parts) if len(parts) > 1 else parts[0]
    return numbers


def generate_plan_docx(
    *,
    company_name: str,
    plan_title: str,
    plan_type: str = "comprehensive",
    plan_number: str = "",
    version_number: str = "",
    sections: list[dict],
    signers: list[dict] | None = None,
    enterprise_info: dict | None = None,
    mermaid_pngs: dict | None = None,
    quality_evidence: dict | None = None,
) -> Document:
    """生成完整的应急预案 DOCX 文档。

    Args:
        company_name: 企业名称
        plan_title: 预案标题
        plan_type: 预案类型 (comprehensive/special/onsite)
        plan_number: 预案编号
        version_number: 版本号
        sections: 章节列表 [{"title": "总则", "level": 1, "content": "..."}]
        signers: 签署人列表 [{"seq": 1, "name": "...", "title": "..."}]
        enterprise_info: 企业附加信息
        mermaid_pngs: Mermaid 流程图 PNG 字节缓存 {hash: bytes}
    """
    import builtins
    builtins.print("!!! generate_plan_docx CALLED !!!", flush=True)
    doc = Document()

    # 1) 注册所有样式
    register_all_styles(doc)

    # 2) 第一节：封面 + 批准页（宽边距）
    first_section = doc.sections[0]
    set_page_margins(first_section,
                     MARGIN_COVER_LEFT, MARGIN_COVER_RIGHT,
                     MARGIN_COVER_TOP, MARGIN_COVER_BOTTOM)

    type_names = {"comprehensive": "综合应急预案", "special": "专项应急预案", "onsite": "现场处置方案"}
    body_title = type_names.get(plan_type, "应急预案")

    build_cover(doc,
                plan_number=plan_number,
                version_number=version_number,
                company_name=company_name,
                doc_title="生产安全事故应急预案",
                signature_company=company_name)

    # 3) 签署页
    if signers:
        build_signature_page(doc, signers)

    # 4) 正文节（标准边距）
    add_section(doc,
                MARGIN_BODY_LEFT, MARGIN_BODY_RIGHT,
                MARGIN_BODY_TOP, MARGIN_BODY_BOTTOM)

    # 正文大标题
    add_body_title(doc, body_title)

    # 5) 逐章节写入
    sec_numbers = _build_docx_section_numbers(sections)
    for idx, section in enumerate(sections):
        title = section.get("title", "")
        level = section.get("level", 1)
        content = section.get("content", "")
        mermaid_svgs = section.get("mermaid_svgs") or {}

        if not title:
            continue

        # 写标题（含编号）
        section_level = level
        heading_level = min(level + 1, 6)
        num = sec_numbers.get(idx, "")
        numbered_title = f"{num} {title}" if num else title
        add_heading(doc, numbered_title, heading_level)

        # 质量证据高亮：正文中出现的片段包 <span background-color:yellow>（_add_inline_runs 转 Word 高亮）
        for ev in (quality_evidence or {}).get(section.get("section_key", ""), []):
            if ev and ev in content:
                content = content.replace(ev, f'<span style="background-color: yellow">{ev}</span>')
        if not content or not content.strip():
            continue

        # 预处理 Mermaid 代码 — 仅当内容不含已被包裹的 Mermaid 时才包装
        if '<code class="language-mermaid"' not in content and '```mermaid' not in content:
            content = _wrap_raw_mermaid(content)

        # Markdown → HTML（始终转换，md_in_html 处理 HTML 内嵌 Markdown）
        import markdown
        content = markdown.markdown(content, extensions=["tables", "fenced_code", "md_in_html"])

        # 修复 Markdown 表格
        content = fix_markdown_tables(content)

        # 提取 Mermaid 代码
        from app.services.mermaid_renderer import _extract_mermaid_code, replace_mermaid_with_placeholders
        codes = _extract_mermaid_code(content)
        cleaned, _ = replace_mermaid_with_placeholders(content)

        # 收集所有需要渲染的 Mermaid 图表（去重）
        logger.info('generate_plan_docx: section=%s, codes=%d, svgs_in_cache=%d',
                     title, len(codes), len(mermaid_svgs))
        # 来源1：<code class="language-mermaid"> （旧格式，直接从 Markdown 转换来的）
        # 来源2：<div class="mermaid-rendered" data-mermaid-hash="..."> （新格式，被 _embed_mermaid_svgs 替换后）
        # 来源3：mermaid_svgs 缓存（内容中无水印代码但缓存有 SVG）
        rendered_hashes = set()
        mermaid_items = []  # (hash, code_or_none, svg_or_none)
        import re as _re2

        # 来源1: <code class="language-mermaid"> → 仅作为去重标记，不独立生成项
        # 这些代码块的 SVG 由来源2(mermaid-rendered div 内联)或来源3(mermaid_svgs 缓存)覆盖
        for code in codes:
            h = _mermaid_hash(code)
            rendered_hashes.add(h)

        # 来源2: <div class="mermaid-rendered" data-mermaid-hash="..."><svg>...</svg></div>
        # 提取并移除这些块
        _mermaid_div_pattern = r'<div class="mermaid-rendered"[^>]*data-mermaid-hash="([^"]+)"[^>]*>(<svg[^>]*>.*?</svg>)</div>'
        for _h, _svg in _re2.findall(_mermaid_div_pattern, cleaned, _re2.DOTALL):
            if _h not in rendered_hashes:
                mermaid_items.append((_h, None, _svg))
                rendered_hashes.add(_h)
        # 从内容中移除所有 mermaid-rendered div
        cleaned = _re2.sub(r'<div class="mermaid-rendered"[^>]*>.*?</div>', '', cleaned, flags=_re2.DOTALL)

        # 来源3: mermaid_svgs 中尚未覆盖的条目
        for h, svg in mermaid_svgs.items():
            if h not in rendered_hashes:
                mermaid_items.append((h, None, svg))
                rendered_hashes.add(h)

        logger.info('generate_plan_docx: total mermaid_items to render=%d', len(mermaid_items))

        # 转换为 DOCX 内容
        html_to_docx_content(doc, cleaned, base_level=heading_level)

                # 插入 Mermaid 图片（sync Playwright，线程内直接调用）
        import base64 as _b64, os as _os
        import html as _html
        from playwright.sync_api import sync_playwright

        _mermaid_js_path = _os.path.join(_os.path.dirname(__file__), "mermaid.min.js")
        _mermaid_js_content = ""
        if _os.path.exists(_mermaid_js_path):
            with open(_mermaid_js_path, "r", encoding="utf-8") as _mf:
                _mermaid_js_content = _mf.read()

        _browser = None
        try:
            _pw = sync_playwright().start()
            _browser = _pw.chromium.launch(headless=True, args=["--no-sandbox", "--disable-setuid-sandbox"])
        except Exception as _pwe:
            logger.warning("Playwright init failed: %s", _pwe)

        for h, code, svg_raw in mermaid_items:
            png_bytes = None
            try:
                if svg_raw and _browser:
                    _page = _browser.new_page(viewport={"width": 900, "height": 600})
                    try:
                        _page.set_content("<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>" + svg_raw + "</body></html>", wait_until="domcontentloaded", timeout=10000)
                        _page.wait_for_selector("svg", timeout=5000)
                        _page.wait_for_timeout(200)
                        _el = _page.query_selector("svg")
                        if _el:
                            png_bytes = _el.screenshot(type="png")
                    finally:
                        _page.close()
                    if png_bytes:
                        logger.info("Mermaid SVG->PNG rendered %s (%d bytes)", h, len(png_bytes))

                if png_bytes is None and code and _browser and _mermaid_js_content:
                    _page = _browser.new_page(viewport={"width": 900, "height": 600})
                    try:
                        _html_str = '<!DOCTYPE html><html><head><meta charset="utf-8"></head><body><div class="mermaid">' + _html.escape(code) + '</div><div id="status"></div><script>' + _mermaid_js_content + '</script><script>mermaid.initialize({startOnLoad:false,theme:"default",securityLevel:"loose"});mermaid.run().then(function(){var svg=document.querySelector(".mermaid svg");if(svg)document.getElementById("status").innerHTML=svg.outerHTML;});</script></body></html>'
                        _page.set_content(_html_str, wait_until="domcontentloaded", timeout=8000)
                        _page.wait_for_selector("#status svg", timeout=8000)
                        _page.wait_for_timeout(300)
                        _el = _page.query_selector("#status svg")
                        if _el:
                            png_bytes = _el.screenshot(type="png")
                    finally:
                        _page.close()
                    if png_bytes:
                        logger.info("Mermaid code->PNG rendered %s (%d bytes)", h, len(png_bytes))
            except Exception as fe:
                logger.warning("Mermaid render %s failed: %s", h, fe)

            if png_bytes:
                img_stream = io.BytesIO(png_bytes)
                doc.add_picture(img_stream, width=Cm(14.6))
                if doc.paragraphs:
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 附图：非占位 SVG → PNG 插入；占位 → 文字行
        for key, meta in (section.get("diagram_svgs") or {}).items():
            if not isinstance(meta, dict):
                continue
            if meta.get("placeholder"):
                add_normal_paragraph(doc, f"【{key}】待补充企业数据后生成")
                continue
            svg = meta.get("svg")
            if not svg:
                continue
            png_bytes = None
            if _browser:
                try:
                    _page = _browser.new_page(viewport={"width": 900, "height": 600})
                    try:
                        _page.set_content(
                            "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>" + svg + "</body></html>",
                            wait_until="domcontentloaded", timeout=10000,
                        )
                        _page.wait_for_selector("svg", timeout=5000)
                        _page.wait_for_timeout(200)
                        _el = _page.query_selector("svg")
                        if _el:
                            png_bytes = _el.screenshot(type="png")
                    finally:
                        _page.close()
                except Exception as e:
                    logger.warning("Diagram %s SVG->PNG failed: %s", key, e)
            if png_bytes:
                img_stream = io.BytesIO(png_bytes)
                doc.add_picture(img_stream, width=Cm(14.6))
                if doc.paragraphs:
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_normal_paragraph(doc, f"【{key}】附图渲染失败")

        if _browser:
            _browser.close()
            _pw.stop()

        # 仅大章节(level==1)之间分页，小节(level>=2)连续，末尾不换页
        if section_level == 1:
            doc.add_page_break()

    return doc


def _mermaid_hash(code: str) -> str:
    """Mermaid 代码哈希 - 与 mermaid_renderer 保持一致，使用 SHA256。"""
    import hashlib
    return hashlib.sha256(code.encode('utf-8')).hexdigest()[:16]


# ── 从 export.py 借用的辅助函数 ──

_MERMAID_KEYWORDS = [
    "flowchart ", "graph ", "sequenceDiagram", "classDiagram",
    "stateDiagram", "erDiagram", "gantt", "pie", "gitGraph",
    "mindmap", "timeline", "journey", "quadrantChart",
]


def _wrap_raw_mermaid(content: str) -> str:
    """检测未包裹的 Mermaid 代码并添加代码围栏。"""
    lines = content.split("\n")
    result = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        is_start = any(line.startswith(kw) for kw in _MERMAID_KEYWORDS)
        is_edge = ("-->" in line or " -- " in line) and not line.startswith("#") and not line.startswith("|")

        if is_start:
            mermaid_lines = [lines[i]]
            j = i + 1
            while j < len(lines):
                nl = lines[j].rstrip()
                if nl.strip() == "":
                    mermaid_lines.append(lines[j])
                    j += 1
                    continue
                is_ml = (
                    "-->" in nl.strip() or " -- " in nl.strip() or
                    (nl.strip()[0:1] in "ABCDEFGHIJKLMNOPQRSTUVWXYZ" and
                     ("[" in nl.strip() or "(" in nl.strip() or "{" in nl.strip()))
                ) or nl.strip().startswith("subgraph") or nl.strip().startswith("end")
                if not is_ml:
                    break
                mermaid_lines.append(lines[j])
                j += 1
            code = "\n".join(mermaid_lines).strip()
            if code:
                result.append("```mermaid")
                result.append(code)
                result.append("```")
            i = j
        elif is_edge:
            mermaid_lines = [lines[i]]
            j = i + 1
            while j < len(lines):
                nl = lines[j].rstrip()
                if nl.strip() == "":
                    mermaid_lines.append(lines[j])
                    j += 1
                    continue
                is_ml = (
                    "-->" in nl.strip() or " -- " in nl.strip() or
                    (nl.strip()[0:1] in "ABCDEFGHIJKLMNOPQRSTUVWXYZ" and
                     ("[" in nl.strip() or "(" in nl.strip() or "{" in nl.strip()))
                ) or nl.strip().startswith("subgraph") or nl.strip().startswith("end")
                if not is_ml:
                    break
                mermaid_lines.append(lines[j])
                j += 1
            code = "\n".join(mermaid_lines).strip()
            if code:
                result.append("```mermaid")
                result.append(code)
                result.append("```")
            i = j
        else:
            result.append(lines[i])
            i += 1
    return "\n".join(result)
