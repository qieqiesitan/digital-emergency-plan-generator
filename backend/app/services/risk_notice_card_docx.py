"""风险告知卡 Word 渲染：A4 竖版、每卡一页、右上角二维码、左栏键值表格 + 安全标志 PNG。

导出端点在 async 上下文中调用：SVG→PNG 复用 mermaid_renderer 的 Playwright 通道
（不可用时回退占位 PNG，不阻断导出）；docx 布局按设计规格 §4/§11 实现。
"""

import io
import logging
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

from app.schemas.risk_notice_card import CardData

logger = logging.getLogger(__name__)

SIGNS_DIR = Path(__file__).resolve().parent.parent / "static" / "signs"

FONT_HEITI = "黑体"

_PNG_HEADER = b"\x89PNG\r\n\x1a\n"


def make_qr_png(url: str) -> bytes:
    """生成二维码 PNG bytes（内容为公开页完整 URL）。"""
    img = qrcode.make(url)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _placeholder_png() -> bytes:
    """生成合法 1x1 PNG 占位图（SVG→PNG 失败时兜底，保证导出不中断）。"""
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (1, 1), color="white").save(buf, format="PNG")
    return buf.getvalue()


async def svg_to_png(svg_name: str) -> bytes:
    """SVG 标志 → PNG bytes；文件缺失或转换失败返回占位 PNG（不阻断导出）。"""
    svg_path = SIGNS_DIR / f"{svg_name}.svg"
    if not svg_path.is_file():
        logger.warning("安全标志 SVG 不存在: %s", svg_name)
        return _placeholder_png()
    try:
        from app.services.mermaid_renderer import render_svg_to_png

        return await render_svg_to_png(svg_path.read_text(encoding="utf-8"))
    except Exception:
        logger.warning("安全标志 SVG→PNG 转换失败: %s，使用占位图", svg_name)
        return _placeholder_png()


def _set_run(run, text: str, size: float = 10, bold: bool = False, color: str | None = None):
    """设置 run 文本与字体（含东亚字体），颜色支持 '#RRGGBB'。"""
    run.text = text
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_HEITI
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_HEITI)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def _shade_paragraph(paragraph, fill_hex: str):
    """为段落设置整行底纹（用于色带/深色标题条）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = p_pr.makeelement(qn("w:shd"), {})
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))


def _shade_cell(cell, fill_hex: str):
    """为表格单元格设置背景色（标签列灰底）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = tc_pr.makeelement(qn("w:shd"), {})
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))


def _render_header(doc, card: CardData):
    """头部三区：企业名（左）/ 居中标题 / 右上角二维码（约 1.4cm）。"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_left, cell_mid, cell_right = table.rows[0].cells

    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_left.add_run(), card.enterprise_name, size=9, color="666666")

    p_mid = cell_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p_mid.add_run(), f"{card.name}安全风险告知卡", size=15, bold=True)

    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_run = p_right.add_run()
    qr_run.add_picture(io.BytesIO(make_qr_png(card.public_url)), width=Cm(1.4))


def _render_level_band(doc, card: CardData):
    """等级色带：全宽底纹 + 白字「{等级}风险」。"""
    p = doc.add_paragraph()
    _shade_paragraph(p, card.level_color or "#d9d9d9")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _set_run(p.add_run(), f"{card.level}风险", size=12, bold=True, color="FFFFFF")


def _render_left(doc, card: CardData, sign_pngs: dict[str, bytes]):
    """左栏：等级色带 + 键值表格 6 行 + 「安全标志」标题 + 标志 PNG。"""
    _render_level_band(doc, card)

    rows = [
        ("风险点名称", card.name),
        ("风险点编号", card.code),
        ("风险等级", card.level),
        ("责任单位", card.responsible_unit),
        ("责任人", card.responsible_person),
        ("联系电话", card.contact_phone),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, value) in enumerate(rows):
        cell_key, cell_value = table.rows[i].cells
        p_key = cell_key.paragraphs[0]
        _set_run(p_key.add_run(), key, size=10, bold=True)
        _shade_cell(cell_key, "F2F2F2")
        p_value = cell_value.paragraphs[0]
        _set_run(p_value.add_run(), value, size=10, bold=True)

    p_title = doc.add_paragraph()
    _shade_paragraph(p_title, "434343")
    p_title.paragraph_format.space_before = Pt(6)
    _set_run(p_title.add_run(), "安全标志", size=11, bold=True, color="FFFFFF")

    sign_par = doc.add_paragraph()
    for sign in card.signs:
        png = sign_pngs.get(sign.svg_name) if sign_pngs else None
        if png and png[:8] == _PNG_HEADER:
            pic_run = sign_par.add_run()
            pic_run.add_picture(io.BytesIO(png), width=Cm(1.5))
        _set_run(sign_par.add_run(), f" {sign.name}  ", size=9)


def _render_right(doc, card: CardData):
    """右栏：四个信息块，标题深色底 + 白字加粗，正文小字。"""
    accident_text = "、".join(card.accident_types)
    accident_text = f"{accident_text}（GB 6441 事故类别）" if accident_text else ""
    blocks = [
        ("主要危险因素描述", card.hazard_description),
        ("主要事故类型", accident_text),
        ("主要风险控制措施", "\n".join(card.control_measures)),
        ("应急处置措施", "\n".join(card.emergency_measures)),
    ]
    for title, body in blocks:
        p_title = doc.add_paragraph()
        _shade_paragraph(p_title, "434343")
        p_title.paragraph_format.space_before = Pt(6)
        _set_run(p_title.add_run(), title, size=11, bold=True, color="FFFFFF")
        body = body.strip() if body else ""
        if not body:
            body = "暂无，请先完善风险评估数据"
        for line in body.split("\n"):
            p_body = doc.add_paragraph()
            _set_run(p_body.add_run(), line, size=10)


def _render_footer(doc, card: CardData):
    """页脚：签发单位（企业名）/ 编制日期 / 版本（V1.0 规则基线 / V1.{version} 快照）。"""
    if card.snapshot and card.snapshot.get("version"):
        version = f"V1.{card.snapshot['version']}"
    else:
        version = "V1.0"
    date = (card.generated_at or "")[:10]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    _set_run(
        p.add_run(),
        f"签发单位：{card.enterprise_name}　|　编制日期：{date}　|　版本：{version}",
        size=8,
        color="666666",
    )


def render_cards_docx(
    cards: list[CardData],
    out_path: str,
    sign_pngs: dict[str, bytes] | None = None,
):
    """渲染多张风险告知卡为 A4 竖版 docx，每卡一页，卡间分页符。"""
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = section.bottom_margin = Cm(1.5)
    sign_pngs = sign_pngs or {}

    for i, card in enumerate(cards):
        _render_header(doc, card)
        _render_left(doc, card, sign_pngs)
        _render_right(doc, card)
        _render_footer(doc, card)
        if i < len(cards) - 1:
            doc.add_page_break()
    doc.save(out_path)
