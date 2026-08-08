import json, os, re, logging
from bs4 import BeautifulSoup
from datetime import datetime, timezone
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy import select
from sse_starlette.sse import EventSourceResponse
from app.database import get_db, async_session
from app.dependencies import get_current_user
from app.models.enterprise import Enterprise, AIConfig
from app.models.risk_assessment import RiskAssessmentReport
from app.schemas.risk_assessment import (
    RiskAssessmentGenerateRequest,
    RiskAssessmentReportResponse,
    RiskAssessmentPreviewResponse,
)
from app.schemas.common import ApiResponse
from app.services.llm_client import decrypt_api_key, llm_chat_completion, llm_stream_all, LLMError
from app.services.markdown_utils import md_to_html
from app.services.sse_utils import sse_event
from app.services.risk_assessment_service import (
    CHAPTER_DEFINITIONS as RA_CHAPTER_DEFINITIONS,
)
from app.services.risk_context_builder import build_risk_management_context
from app.services.risk_assessment_service import (
    build_chapter_prompt,
    get_chapter_keys,
    get_chapter_title,
    _get_ra_system_prompt,
)
from app.config import settings

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/enterprises", tags=["Risk Assessment"])


def _html_table_to_docx(doc, html_table: str):
    """将 HTML <table> 渲染为 python-docx 表格"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    soup = BeautifulSoup(html_table, "html.parser")
    table_el = soup.find("table")
    if not table_el:
        return
    rows = table_el.find_all("tr")
    if not rows:
        return
    # Determine column count from first row
    first_cells = rows[0].find_all(["th", "td"])
    col_count = len(first_cells)
    if col_count == 0:
        return
    docx_table = doc.add_table(rows=len(rows), cols=col_count)
    docx_table.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for ci, cell in enumerate(cells):
            if ci >= col_count:
                break
            docx_cell = docx_table.cell(ri, ci)
            text = cell.get_text(strip=True)
            docx_cell.text = text
            # Bold for header cells
            if cell.name == "th":
                for p in docx_cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            else:
                for p in docx_cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    doc.add_paragraph("")  # spacing after table


def _clean_for_docx(content: str) -> str:
    """Strip markdown artifacts that _render_content_to_docx doesn't handle."""
    import re as _re
    # 1. Remove fenced code blocks
    content = _re.sub(r'```[\w]*\n[\s\S]*?```', '', content)
    # 2. Remove bare mermaid blocks (no code fences: "mermaid\nflowchart...")
    content = _re.sub(r'\nmermaid\n(?:flowchart|graph|sequenceDiagram|pie|mindmap|classDiagram|stateDiagram|erDiagram|gantt|journey|gitgraph)[\s\S]*?(?=\n\n[^A-Za-z\-\[]>]|\n(?:json|```)\n|\Z)', '', content)
    # 3. Remove bare "json\n{...}\n```" blocks (opening ``` missing, closing present)
    content = _re.sub(r'\njson\n\{[\s\S]*?\n```', '', content)
    # 4. Strip trailing ``` backticks before JSON detection
    content = _re.sub(r'\n```\s*$', '', content)
    # 5. Remove trailing JSON summary block (handles nested braces)
    def _strip_trailing_json(s: str) -> str:
        """Find and remove a trailing JSON object with balanced braces."""
        stripped = s.rstrip()
        if not stripped.endswith('}'):
            return s
        # Count braces backwards
        depth = 0
        for i in range(len(stripped) - 1, -1, -1):
            ch = stripped[i]
            if ch == '}':
                depth += 1
            elif ch == '{':
                depth -= 1
                if depth == 0:
                    prefix = stripped[:i].rstrip()
                    if prefix:
                        return prefix + '\n'
                    return prefix
        return s
    content = _strip_trailing_json(content)
    # 3. Strip **bold** markers
    content = content.replace('**', '')
    # 4. Convert * list markers to - (already handled)
    content = _re.sub(r'^(\s*)\* ', r'\1- ', content, flags=_re.MULTILINE)
    # 5. Fix #text -> # text (missing space after #)
    content = _re.sub(r'^(#{1,6})([^\s#])', r'\1 \2', content, flags=_re.MULTILINE)
    # 6. Remove duplicate lines immediately following a heading
    lines_out = []
    prev_line = None
    for line in content.split('\n'):
        stripped = line.strip()
        # If current line is a duplicate of the previous line's content (after heading prefix)
        if prev_line and stripped and stripped == prev_line:
            continue
        lines_out.append(line)
        # Extract plain text from heading for next-line dedup
        m = _re.match(r'^#{1,6}\s+(.+)', stripped)
        if m:
            prev_line = m.group(1).strip()
        elif stripped:
            prev_line = stripped
        else:
            prev_line = None
    content = '\n'.join(lines_out)
    return content

def _render_content_to_docx(doc, content: str):
    """Render content that may contain interleaved text and HTML tables into docx."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    content = _clean_for_docx(content)
    # Split content by HTML table blocks
    parts = re.split(r"(<table[\s\S]*?</table>)", content)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.startswith("<table"):
            _html_table_to_docx(doc, part)
        else:
            # Render text lines
            for line in part.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("# "):
                    h = doc.add_heading(level=1)
                    h.add_run(line[2:])
                elif line.startswith("## "):
                    h = doc.add_heading(level=2)
                    h.add_run(line[3:])
                elif line.startswith("### "):
                    h = doc.add_heading(level=3)
                    h.add_run(line[4:])
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif re.match(r"^\d+[.)] ", line):
                    doc.add_paragraph(re.sub(r"^\d+[.)]\s*", "", line), style="List Number")
                elif re.match(r"^\d+）", line):
                    doc.add_paragraph(re.sub(r"^\d+）\s*", "", line), style="List Number")
                else:
                    doc.add_paragraph(line)

# ---- LLM streaming utilities (used by both risk_assessment and resource_investigation) ----

def _guard_decrypt(ai_config) -> None:
    """保持原 decrypt 失败 → HTTPException(500) 的语义。"""
    try:
        decrypt_api_key(ai_config.api_key_encrypted)
    except Exception:
        raise HTTPException(500, "AI config key decryption failed")


async def _stream_llm_with_messages(messages: list[dict], ai_config: AIConfig) -> str:
    _guard_decrypt(ai_config)
    try:
        return await llm_stream_all(messages, ai_config, timeout=120)
    except LLMError as e:
        # 保持原 risk_assessment 文案
        raise Exception(f"LLM call failed: {e.status_code} {e.text[:300]}")


async def _stream_llm_with_messages_chunked(messages: list[dict], ai_config: AIConfig):
    _guard_decrypt(ai_config)
    try:
        gen = await llm_chat_completion(messages, ai_config, stream=True, timeout=120)
        async for chunk in gen:
            yield chunk
    except LLMError as e:
        raise Exception(f"LLM call failed: {e.status_code} {e.text[:300]}")


async def _stream_llm_with_system(prompt: str, ai_config: AIConfig) -> str:
    messages = [
        {"role": "system", "content": _get_ra_system_prompt()},
        {"role": "user", "content": prompt},
    ]
    return await _stream_llm_with_messages(messages, ai_config)


# ---- API Endpoints ----

@router.get("/{enterprise_id}/risk-assessment")
async def get_risk_assessment(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (await db.execute(
        select(Enterprise).where(Enterprise.id == enterprise_id, Enterprise.user_id == current_user.id)
    )).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "企业不存在")

    report = (await db.execute(
        select(RiskAssessmentReport).where(
            RiskAssessmentReport.enterprise_id == enterprise_id,
            RiskAssessmentReport.status.in_(["completed", "draft"]),
        )
    )).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到已完成的风险评估报告")

    return ApiResponse(data=RiskAssessmentReportResponse.model_validate(report))


@router.get("/{enterprise_id}/risk-assessment/summary")
async def get_risk_assessment_summary(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (await db.execute(
        select(Enterprise).where(Enterprise.id == enterprise_id, Enterprise.user_id == current_user.id)
    )).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "企业不存在")

    report = (await db.execute(
        select(RiskAssessmentReport).where(
            RiskAssessmentReport.enterprise_id == enterprise_id,
            RiskAssessmentReport.status.in_(["completed", "draft"]),
        )
    )).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到已完成的风险评估报告")

    return ApiResponse(data=report.summary or {})


@router.get("/{enterprise_id}/risk-assessment/preview")
async def preview_risk_assessment(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (await db.execute(
        select(Enterprise).where(Enterprise.id == enterprise_id, Enterprise.user_id == current_user.id)
    )).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "企业不存在")

    report = (await db.execute(
        select(RiskAssessmentReport).where(
            RiskAssessmentReport.enterprise_id == enterprise_id,
            RiskAssessmentReport.status.in_(["completed", "draft"]),
        )
    )).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到已完成的风险评估报告")

    html = md_to_html(_clean_for_docx(report.content), output_format="html5")
    return ApiResponse(data=RiskAssessmentPreviewResponse(
        report_id=report.id, title=report.title, html=html
    ))


@router.get("/{enterprise_id}/risk-assessment/export")
async def export_risk_assessment(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (await db.execute(
        select(Enterprise).where(Enterprise.id == enterprise_id, Enterprise.user_id == current_user.id)
    )).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "企业不存在")

    report = (await db.execute(
        select(RiskAssessmentReport).where(
            RiskAssessmentReport.enterprise_id == enterprise_id,
            RiskAssessmentReport.status.in_(["completed", "draft"]),
        )
    )).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到已完成的风险评估报告")

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(500, "python-docx 未安装")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)

    # ---- Professional Cover Page ----
    # Add empty paragraphs for spacing (top margin)
    for _ in range(6):
        doc.add_paragraph("")

    # Main title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(ent.name)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Report type subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("生产安全事故风险评估报告")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Cover footer
    for line in [
        f"编制单位：{ent.name}",
        f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
    ]:
        if line:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    # ---- Report Body ----
    _render_content_to_docx(doc, report.content)

    # ---- Export ----
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    safe_name = ent.name.replace(" ", "_") if ent else "企业"
    filename = f"{safe_name}_事故风险评估报告.docx"
    path = os.path.join(settings.EXPORT_DIR, filename)
    doc.save(path)
    return FileResponse(
        path, filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )



@router.post("/{enterprise_id}/risk-assessment/generate")
async def generate_risk_assessment(
    enterprise_id: str,
    request: RiskAssessmentGenerateRequest = RiskAssessmentGenerateRequest(),
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (await db.execute(
        select(Enterprise).where(Enterprise.id == enterprise_id, Enterprise.user_id == current_user.id)
    )).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "企业不存在")

    context = await build_risk_management_context(enterprise_id, db)
    if context["total_events"] == 0:
        raise HTTPException(400, "请先录入风险分级管控数据")

    ai_config = (await db.execute(
        select(AIConfig).where(AIConfig.user_id == current_user.id)
    )).scalar_one_or_none()
    if not ai_config:
        raise HTTPException(400, "请先配置 AI 模型")

    existing = (await db.execute(
        select(RiskAssessmentReport).where(
            RiskAssessmentReport.enterprise_id == enterprise_id,
            RiskAssessmentReport.status == "generating",
        )
    )).scalar_one_or_none()
    if existing:
        raise HTTPException(400, "已有正在生成的报告，请等待完成")

    report = (await db.execute(
        select(RiskAssessmentReport).where(RiskAssessmentReport.enterprise_id == enterprise_id)
    )).scalar_one_or_none()

    title = f"{ent.name} 事故风险评估报告"
    if report:
        report.title = title
        report.content = ""
        report.summary = {}
        report.status = "generating"
    else:
        report = RiskAssessmentReport(
            enterprise_id=enterprise_id, title=title, status="generating",
        )
        db.add(report)
    await db.commit()

    async def event_generator():
        full_content = ""
        chapter_contents: list[dict] = []
        try:
            chapter_keys = get_chapter_keys()
            total = len(chapter_keys)
            yield sse_event("progress", message=f"开始逐章生成风险评估报告（共{total}章）...",
                       current=0, total=total)

            for i, ck in enumerate(chapter_keys):
                ctitle = get_chapter_title(ck)
                yield sse_event("progress",
                           message=f"正在生成「{ctitle}」（{i+1}/{total}）",
                           current=i+1, total=total, section_key=ck)

                ch_prompt = build_chapter_prompt(
                    ck, context,
                    previous_chapters=chapter_contents if chapter_contents else None,
                    custom_instruction=request.custom_instruction,
                )
                messages = [
                    {"role": "system", "content": _get_ra_system_prompt()},
                    {"role": "user", "content": ch_prompt},
                ]
                ch_content = ""
                async for chunk_content in _stream_llm_with_messages_chunked(messages, ai_config):
                    ch_content += chunk_content
                    yield sse_event("chunk", content=chunk_content, section_key=ck)

                chapter_contents.append({
                    "key": ck, "title": ctitle, "content": ch_content,
                })
                full_content += f"\n\n{ctitle}\n\n{ch_content}"
                yield sse_event("section_done", section_key=ck,
                           message=f"「{ctitle}」生成完成",
                           completed=i+1, total=total)

            # Save chapter contents to summary, set status to draft (user will merge manually)
            chapters_json = [
                {"key": c["key"], "title": c["title"], "content": c["content"]}
                for c in chapter_contents
            ]

            async with async_session() as bg_db:
                bg_report = (await bg_db.execute(
                    select(RiskAssessmentReport).where(RiskAssessmentReport.id == report.id)
                )).scalar_one_or_none()
                if bg_report:
                    bg_report.status = "draft"
                    bg_report.content = full_content.strip()
                    bg_report.summary = {"chapters": chapters_json}
                    try:
                        import json as _json2, re as _re2
                        last_ch = chapter_contents[-1] if chapter_contents else None
                        if last_ch:
                            m = _re2.search(r"\{[^}]+\}\s*$", last_ch.get("content", ""))
                            if m:
                                struct = _json2.loads(m.group())
                                bg_report.summary.update(struct)
                    except Exception:
                        pass
                    await bg_db.commit()

            import json as _json
            yield sse_event("batch_done", report_id=report.id,
                       message=f"报告生成完成，共{total}章",
                       completed=total, total=total,
                       chapters=_json.dumps(chapters_json, ensure_ascii=False))
        except Exception as e:
            import traceback
            logger.error(f"Risk assessment generation failed: {e}\n{traceback.format_exc()}")
            async with async_session() as bg_db:
                bg_report = (await bg_db.execute(
                    select(RiskAssessmentReport).where(RiskAssessmentReport.id == report.id)
                )).scalar_one_or_none()
                if bg_report:
                    bg_report.status = "draft"
                    bg_report.content = full_content
                    await bg_db.commit()
            yield sse_event("error", message=str(e))
    return EventSourceResponse(event_generator())


@router.get("/{enterprise_id}/risk-assessment/chapters")
async def get_risk_assessment_chapters():
    """Return chapter definitions for risk assessment report (used by frontend)."""
    return ApiResponse(data=[
        {"key": c["key"], "title": c["title"]}
        for c in RA_CHAPTER_DEFINITIONS
    ])


@router.post("/{enterprise_id}/risk-assessment/merge")
async def merge_risk_assessment(
    enterprise_id: str,
    request: RiskAssessmentGenerateRequest = RiskAssessmentGenerateRequest(),
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    """Merge edited chapters into final risk assessment report."""
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "ERROR")

    import json as _json
    chapters_data = request.custom_instruction

    report = (
        await db.execute(
            select(RiskAssessmentReport).where(
                RiskAssessmentReport.enterprise_id == enterprise_id
            )
        )
    ).scalar_one_or_none()

    if not report:
        report = RiskAssessmentReport(
            enterprise_id=enterprise_id,
            title="",
            status="draft",
        )
        db.add(report)
        await db.commit()

    chapters = []
    try:
        if chapters_data:
            chapters = _json.loads(chapters_data)
    except Exception:
        raise HTTPException(400, "ERROR")

    if not chapters:
        raise HTTPException(400, "ERROR")

    report_title = f"#{ent.name} 生产安全事故风险评估报告"
    merged_parts = []
    for ch in chapters:
        merged_parts.append(f"## {ch.get('title', '')}\n\n{ch.get('content', '')}")
    merged = report_title + "\n\n" + "\n\n".join(merged_parts)

    report.title = report_title
    merged = _clean_for_docx(merged)
    report.content = merged
    report.status = "completed"
    report.summary = {"chapters": chapters}
    try:
        import json as _json2, re as _re2
        last_ch = chapters[-1] if chapters else None
        if last_ch:
            m = _re2.search(r"\{[^}]+\}\s*$", last_ch.get("content", ""))
            if m:
                struct = _json2.loads(m.group())
                report.summary.update(struct)
    except Exception:
        pass
    report.generated_at = datetime.now(timezone.utc)
    await db.commit()

    return ApiResponse(data={"report_id": report.id, "title": report_title, "status": "completed"})
