import json, os, markdown, logging
from datetime import datetime, timezone
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sse_starlette.sse import EventSourceResponse
from app.database import get_db, async_session
from app.dependencies import get_current_user
from app.models.enterprise import Enterprise, EmergencyResource, AIConfig
from app.models.resource_investigation import ResourceInvestigationReport
from app.schemas.resource_investigation import (
    ResourceInvestigationGenerateRequest,
    ResourceInvestigationReportResponse,
    ResourceInvestigationPreviewResponse,
)
from app.schemas.common import ApiResponse
from app.services.resource_investigation_service import (
    CHAPTER_DEFINITIONS as RI_CHAPTER_DEFINITIONS,
    build_resource_investigation_context,
    build_chapter_prompt,
    get_chapter_keys,
    get_chapter_title,
    _get_ri_system_prompt,
)
from app.config import settings
from app.routers.risk_assessment import _stream_llm_with_system, _stream_llm_with_messages_chunked

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/enterprises", tags=["Resource Investigation"])


def _md_to_html(text: str) -> str:
    if not text or text.strip().startswith("<"):
        return text
    return markdown.markdown(text, extensions=["tables", "fenced_code"])




def _sse(event_type: str, **kwargs) -> str:
    obj = {"type": event_type, **kwargs}
    return json.dumps(obj, ensure_ascii=False)


@router.get("/{enterprise_id}/resource-investigation")
async def get_resource_investigation(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "未找到报告")

    report = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id,
                ResourceInvestigationReport.status.in_(["completed", "draft"]),
            )
        )
    ).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到报告")

    return ApiResponse(data=ResourceInvestigationReportResponse.model_validate(report))


@router.get("/{enterprise_id}/resource-investigation/summary")
async def get_resource_investigation_summary(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "未找到报告")

    report = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id,
                ResourceInvestigationReport.status.in_(["completed", "draft"]),
            )
        )
    ).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到报告")

    return ApiResponse(data=report.summary or {})


@router.get("/{enterprise_id}/resource-investigation/preview")
async def preview_resource_investigation(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "未找到报告")

    report = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id,
                ResourceInvestigationReport.status.in_(["completed", "draft"]),
            )
        )
    ).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到报告")

    html = _md_to_html(report.content)
    return ApiResponse(
        data=ResourceInvestigationPreviewResponse(
            report_id=report.id,
            title=report.title,
            html=html,
        )
    )


@router.get("/{enterprise_id}/resource-investigation/export")
async def export_resource_investigation(
    enterprise_id: str,
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "未找到报告")

    report = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id,
                ResourceInvestigationReport.status.in_(["completed", "draft"]),
            )
        )
    ).scalar_one_or_none()
    if not report:
        raise HTTPException(404, "未找到报告")

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "ERROR")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)

    doc.add_paragraph("")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(report.title)
    run.font.size = Pt(22)
    run.bold = True

    if ent:
        for line in [
            f"编制单位：{ent.name}",
            f"Date: {datetime.now().strftime('%Y%m%d')}",
        ]:
            if line:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(line)
    doc.add_page_break()

    for line in report.content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            h = doc.add_heading(level=1)
            h.add_run(line[2:])
        elif line.startswith("## "):
            h = doc.add_heading(level=2)
            h.add_run(line[3:])
        elif line.startswith("### "):
            h = doc.add_heading(level=3)
            h.add_run(line[4:])
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("1. "):
            doc.add_paragraph(line[3:], style="List Number")
        else:
            doc.add_paragraph(line)

    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    safe_name = ent.name.replace(" ", "_") if ent else "企业"
    filename = f"{safe_name}_应急资源调查报告.docx"
    path = os.path.join(settings.EXPORT_DIR, filename)
    doc.save(path)
    return FileResponse(
        path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.post("/{enterprise_id}/resource-investigation/generate")
async def generate_resource_investigation(
    enterprise_id: str,
    request: ResourceInvestigationGenerateRequest = ResourceInvestigationGenerateRequest(),
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "未找到报告")

    # Check resources
    resource_count = (
        await db.execute(
            select(EmergencyResource).where(
                EmergencyResource.enterprise_id == enterprise_id
            )
        )
    ).scalars().all()
    if len(resource_count) == 0:
        raise HTTPException(400, "ERROR")

    # Check AI config
    ai_config = (
        await db.execute(
            select(AIConfig).where(AIConfig.user_id == current_user.id)
        )
    ).scalar_one_or_none()
    if not ai_config:
        raise HTTPException(400, "ERROR")

    # Check for existing generating report
    existing = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id,
                ResourceInvestigationReport.status == "generating",
            )
        )
    ).scalar_one_or_none()
    if existing:
        raise HTTPException(400, "ERROR")

    # Build context
    context = await build_resource_investigation_context(enterprise_id, db)

    # Create or update report record
    report = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id
            )
        )
    ).scalar_one_or_none()

    title = f"{ent.name} 应急资源调查报告"
    if report:
        report.title = title
        report.content = ""
        report.summary = {}
        report.status = "generating"
    else:
        report = ResourceInvestigationReport(
            enterprise_id=enterprise_id,
            title=title,
            status="generating",
        )
        db.add(report)
    await db.commit()

    async def event_generator():
        full_content = ""
        chapter_contents: list[dict] = []
        try:
            chapter_keys = get_chapter_keys()
            total = len(chapter_keys)
            yield _sse("progress", message=f"开始逐章生成应急资源调查报告（共{total}章）...",
                       current=0, total=total)

            for i, ck in enumerate(chapter_keys):
                ctitle = get_chapter_title(ck)
                yield _sse("progress",
                           message=f"正在生成「{ctitle}」（{i+1}/{total}）",
                           current=i+1, total=total, section_key=ck)

                ch_prompt = build_chapter_prompt(
                    ck, context,
                    previous_chapters=chapter_contents if chapter_contents else None,
                    custom_instruction=request.custom_instruction,
                )
                messages = [
                    {"role": "system", "content": _get_ri_system_prompt()},
                    {"role": "user", "content": ch_prompt},
                ]
                ch_content = ""
                async for chunk_content in _stream_llm_with_messages_chunked(messages, ai_config):
                    ch_content += chunk_content
                    yield _sse("chunk", content=chunk_content, section_key=ck)

                chapter_contents.append({
                    "key": ck, "title": ctitle, "content": ch_content,
                })
                full_content += f"\n\n{ctitle}\n\n{ch_content}"
                yield _sse("section_done", section_key=ck,
                           message=f"「{ctitle}」生成完成",
                           completed=i+1, total=total)

            # Save chapter contents to summary, set status to draft (user will merge manually)
            chapters_json = [
                {"key": c["key"], "title": c["title"], "content": c["content"]}
                for c in chapter_contents
            ]

            async with async_session() as bg_db:
                bg_report = (
                    await bg_db.execute(
                        select(ResourceInvestigationReport).where(
                            ResourceInvestigationReport.id == report.id
                        )
                    )
                ).scalar_one_or_none()
                if bg_report:
                    bg_report.status = "draft"
                    bg_report.content = full_content.strip()
                    bg_report.summary = {"chapters": chapters_json}
                    try:
                        import json as _json2, re as _re2
                        last_ch = chapter_contents[-1] if chapter_contents else None
                        if last_ch:
                            m = _re2.search(r"\{[^}]+\}\s*$", last_ch.get("content", ""))
                            if m:
                                struct = _json2.loads(m.group())
                                bg_report.summary.update(struct)
                    except Exception:
                        pass
                    await bg_db.commit()

            import json as _json
            yield _sse("batch_done", report_id=report.id,
                       message=f"报告生成完成，共{total}章",
                       completed=total, total=total,
                       chapters=_json.dumps(chapters_json, ensure_ascii=False))
        except Exception as e:
            import traceback; logger.error(f"Resource investigation generation failed: {e}\n{traceback.format_exc()}")
            async with async_session() as bg_db:
                bg_report = (
                    await bg_db.execute(
                        select(ResourceInvestigationReport).where(
                            ResourceInvestigationReport.id == report.id
                        )
                    )
                ).scalar_one_or_none()
                if bg_report:
                    bg_report.status = "draft"
                    bg_report.content = full_content
                    await bg_db.commit()
            yield _sse("error", message=str(e))

    return EventSourceResponse(event_generator())



@router.get("/{enterprise_id}/resource-investigation/chapters")
async def get_resource_investigation_chapters():
    """Return chapter definitions for resource investigation report (used by frontend)."""
    return ApiResponse(data=[
        {"key": c["key"], "title": c["title"]}
        for c in RI_CHAPTER_DEFINITIONS
    ])


@router.post("/{enterprise_id}/resource-investigation/merge")
async def merge_resource_investigation(
    enterprise_id: str,
    request: ResourceInvestigationGenerateRequest = ResourceInvestigationGenerateRequest(),
    current_user=Depends(get_current_user),
    db=Depends(get_db),
):
    """Merge edited chapters into final report."""
    ent = (
        await db.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.user_id == current_user.id,
            )
        )
    ).scalar_one_or_none()
    if not ent:
        raise HTTPException(404, "未找到报告")

    # Chapters come as JSON string in custom_instruction or as a separate field
    # We accept them via request body extension: chapters field
    import json as _json
    from fastapi import Body
    chapters_data = request.custom_instruction  # The frontend sends chapters here

    report = (
        await db.execute(
            select(ResourceInvestigationReport).where(
                ResourceInvestigationReport.enterprise_id == enterprise_id
            )
        )
    ).scalar_one_or_none()

    if not report:
        raise HTTPException(404, "未找到报告")

    # Parse chapters from the request
    chapters = []
    try:
        if chapters_data:
            chapters = _json.loads(chapters_data)
    except Exception:
        raise HTTPException(400, "ERROR")

    if not chapters:
        raise HTTPException(400, "ERROR")

    # Merge chapters into full report
    report_title = f"#{ent.name} 应急资源调查报告"
    merged_parts = []
    for ch in chapters:
        merged_parts.append(f"## {ch.get('title', '')}\n\n{ch.get('content', '')}")
    merged = report_title + "\n\n" + "\n\n".join(merged_parts)

    report.title = report_title
    report.content = merged
    report.status = "completed"
    report.summary = {"chapters": chapters}
    try:
        import json as _json2, re as _re2
        last_ch = chapters[-1] if chapters else None
        if last_ch:
            m = _re2.search(r"\{[^}]+\}\s*$", last_ch.get("content", ""))
            if m:
                struct = _json2.loads(m.group())
                report.summary.update(struct)
    except Exception:
        pass
    report.generated_at = datetime.now(timezone.utc)
    await db.commit()

    return ApiResponse(data={"report_id": report.id, "title": report_title, "status": "completed"})
