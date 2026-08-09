import io

import pytest

from app.services.file_parser import parse_file_text


def test_parse_csv_text():
    text = parse_file_text("chem.csv", b"name,cas\nmethanol,67-56-1\nethanol,64-17-5\n")
    assert "methanol" in text
    assert "67-56-1" in text


def test_parse_plain_text_txt():
    text = parse_file_text("note.txt", "企业地址：杭州市XX区".encode("utf-8"))
    assert "企业地址" in text


def test_parse_unsupported_extension_raises():
    with pytest.raises(ValueError):
        parse_file_text("file.exe", b"x")


def test_parse_xlsx_bytes():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["化学品名称", "CAS号"])
    ws.append(["甲醇", "67-56-1"])
    buf = io.BytesIO()
    wb.save(buf)
    text = parse_file_text("chem.xlsx", buf.getvalue())
    assert "甲醇" in text
    assert "67-56-1" in text


def test_parse_docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_paragraph("化学品安全技术说明书：甲醇")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "甲醇"
    table.cell(0, 1).text = "67-56-1"
    table.cell(1, 0).text = "乙醇"
    table.cell(1, 1).text = "64-17-5"
    buf = io.BytesIO()
    doc.save(buf)
    text = parse_file_text("chem.docx", buf.getvalue())
    assert "化学品安全技术说明书：甲醇" in text
    assert "67-56-1" in text


def test_parse_pdf_bytes():
    try:
        import fitz  # PyMuPDF
    except ImportError:
        pytest.skip("PyMuPDF 未安装，跳过 PDF 解析测试")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Methanol MSDS test page")
    buf = doc.tobytes()
    doc.close()
    text = parse_file_text("chem.pdf", buf)
    assert "Methanol MSDS test page" in text


def test_parse_corrupt_xlsx_raises_valueerror():
    with pytest.raises(ValueError, match="文件解析失败"):
        parse_file_text("a.xlsx", b"not a zip")


def test_parse_csv_gbk():
    data = "名称,cas\n甲醇,67-56-1".encode("gbk")
    text = parse_file_text("chem.csv", data)
    assert "甲醇" in text
    assert "67-56-1" in text
