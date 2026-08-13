"""风险告知卡 docx 导出测试：二维码 PNG、docx 渲染、导出端点（mock DB）。

渲染单测传入最小占位 PNG；导出端点 200 用例走真实 SVG→PNG（Playwright），
其余用例 mock SVG→PNG 保持快速稳定。
"""
import io
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.config import settings
from app.database import get_db
from app.dependencies import get_current_user
from app.models.enterprise import Enterprise
from app.models.risk_management import RiskEvent, RiskObject
from app.models.user import User
from app.routers import risk_notice_card
from app.schemas.risk_notice_card import CardData, SignItem
from app.services.risk_notice_card_docx import make_qr_png, render_cards_docx


def _placeholder_png() -> bytes:
    """最小合法 PNG（1x1 白色），测试渲染用占位标志。"""
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (1, 1), color="white").save(buf, format="PNG")
    return buf.getvalue()


PLACEHOLDER_PNG = _placeholder_png()


@pytest.fixture(scope="session", autouse=True)
def _close_shared_browser():
    """本模块导出集成测试会拉起 Playwright 共享浏览器，会话结束回收。"""
    yield
    import asyncio

    from app.services import mermaid_renderer

    try:
        asyncio.run(mermaid_renderer._close_browser())
    except Exception:
        pass


def _card(oid: str, name: str, level: str = "重大") -> CardData:
    return CardData(
        object_id=oid,
        enterprise_name="测试公司",
        name=name,
        code="FX-001",
        level=level,
        level_color="#ff4d4f",
        responsible_unit="储运车间",
        responsible_person="张三",
        contact_phone="13800000000",
        fallback_used=False,
        signs=[
            SignItem(
                category="warning", name="当心爆炸", svg_name="warning-explosion"
            )
        ],
        hazard_description="泄漏遇明火引发火灾爆炸",
        accident_types=["火灾", "爆炸"],
        control_measures=["1. 防静电接地", "2. 动火审批"],
        emergency_measures=["1. 切断气源", "2. 报警"],
        public_url="/r/token123",
        generated_at="2026-08-11T00:00:00Z",
    )


def _all_paragraph_texts(doc: Document) -> list[str]:
    """遍历 body 全部段落（含表格单元格内段落）。"""
    return [
        Paragraph(p, doc).text
        for p in doc.element.body.iter(qn("w:p"))
    ]


def _page_breaks(doc: Document) -> int:
    return sum(
        1
        for br in doc.element.body.iter(qn("w:br"))
        if br.get(qn("w:type")) == "page"
    )


def _drawings(doc: Document) -> int:
    return sum(1 for _ in doc.element.body.iter(qn("w:drawing")))


def _embedded_png_blobs(doc: Document) -> list[bytes]:
    """提取 docx 内嵌的全部 PNG 原始字节（含二维码与安全标志）。"""
    blobs: list[bytes] = []
    for drawing in doc.element.body.iter(qn("w:drawing")):
        for blip in drawing.iter(qn("a:blip")):
            rid = blip.get(qn("r:embed"))
            if not rid:
                continue
            part = doc.part.related_parts.get(rid)
            if part is None:
                continue
            blob = part.blob
            if blob[:8] == b"\x89PNG\r\n\x1a\n":
                blobs.append(blob)
    return blobs


def _embedded_png_sizes(doc: Document) -> list[tuple[int, int]]:
    """解析 docx 内嵌 PNG 实际渲染尺寸（IHDR 宽高），区分真实图片与 1x1 占位图。"""
    sizes: list[tuple[int, int]] = []
    for blob in _embedded_png_blobs(doc):
        if len(blob) >= 24:
            width = int.from_bytes(blob[16:20], "big")
            height = int.from_bytes(blob[20:24], "big")
            sizes.append((width, height))
    return sizes


def _decode_qr(png: bytes) -> str:
    """用 OpenCV 解码二维码 PNG，返回内容文本（失败返回空串）。"""
    import cv2
    import numpy as np

    arr = np.frombuffer(png, dtype=np.uint8)
    gray = cv2.imdecode(arr, cv2.IMREAD_GRAYSCALE)
    if gray is None:
        return ""
    text, points, _ = cv2.QRCodeDetector().detectAndDecode(gray)
    return text if points is not None else ""


def _embedded_qr_contents(doc: Document) -> list[str]:
    """解码 docx 内嵌二维码实际内容（空串表示该 PNG 非二维码/解码失败）。"""
    return [t for t in (_decode_qr(b) for b in _embedded_png_blobs(doc)) if t]


def test_make_qr_png_returns_png_bytes():
    png = make_qr_png("http://localhost/r/token123")
    assert png[:8] == b"\x89PNG\r\n\x1a\n"


def test_render_cards_docx_one_page_per_card(tmp_path):
    cards = [_card("o1", "LPG 储罐区"), _card("o2", "配电室", "较大")]
    out = tmp_path / "cards.docx"
    render_cards_docx(
        cards,
        str(out),
        {"warning-explosion": PLACEHOLDER_PNG},
        base_url="http://testserver",
    )
    assert out.exists()

    doc = Document(str(out))
    texts = _all_paragraph_texts(doc)
    titles = [t for t in texts if "安全风险告知卡" in t]
    assert len(titles) == 2  # 每卡一页 → 标题出现 2 次
    assert _page_breaks(doc) == 1  # 2 张卡 → 1 个分页符
    assert any("当心爆炸" in t for t in texts)  # 标志名称文字
    assert any("签发单位：测试公司" in t for t in texts)  # 页脚
    assert any("V1.0" in t for t in texts)  # 规则基线版本
    # 二维码内容必须是完整公开页 URL（相对路径 /r/... 手机扫码无法解析主机）
    assert _embedded_qr_contents(doc) == [
        "http://testserver/r/token123",
        "http://testserver/r/token123",
    ]


def test_render_cards_docx_uses_two_column_body_layout(tmp_path):
    """导出与预览一致：左栏键值/标志，右栏信息块（左右分栏，非上下堆叠）。"""
    cards = [_card("o1", "LPG 储罐区")]
    out = tmp_path / "layout.docx"
    render_cards_docx(
        cards,
        str(out),
        {"warning-explosion": PLACEHOLDER_PNG},
        base_url="http://testserver",
    )
    doc = Document(str(out))
    tables = doc.tables
    assert len(tables) == 2  # header(3列) + body(左右分栏 2列)
    body = tables[1]
    assert len(body.columns) == 2
    left, right = body.rows[0].cells
    # 左栏：嵌套键值表含「风险点名称」；右栏：信息块标题
    left_text = " ".join(c.text for t in left.tables for row in t.rows for c in row.cells)
    right_text = "\n".join(p.text for p in right.paragraphs)
    assert "风险点名称" in left_text
    assert "主要危险因素描述" in right_text
    assert "应急处置措施" in right_text


def test_render_cards_docx_qr_falls_back_to_public_url_without_base(tmp_path):
    """未传 base_url 时（纯函数调用方）二维码保留 public_url 原值，不崩。"""
    out = tmp_path / "cards.docx"
    render_cards_docx(
        [_card("o1", "LPG 储罐区")], str(out), {"warning-explosion": PLACEHOLDER_PNG}
    )
    assert _embedded_qr_contents(Document(str(out))) == ["/r/token123"]


def _enterprise(**overrides):
    ent = Enterprise(
        id="e1",
        user_id="u1",
        name="甲公司",
        safety_officer="李四",
        safety_officer_phone="13900000000",
    )
    for key, value in overrides.items():
        setattr(ent, key, value)
    return ent


def _risk_object(**overrides):
    obj = RiskObject(
        id="o1",
        enterprise_id="e1",
        zone_id="z1",
        name="配电室",
        responsible_unit="动力车间",
        responsible_person="王五",
        contact_phone="13800000000",
        public_token="tok1",
    )
    for key, value in overrides.items():
        setattr(obj, key, value)
    return obj


def _fire_event():
    return RiskEvent(
        accident_type="火灾",
        risk_level="重大",
        trigger_conditions="泄漏遇明火",
        consequences="火灾爆炸",
        method_type="LS",
    )


def _rows_result(rows):
    res = MagicMock()
    res.scalars.return_value.all.return_value = rows
    return res


def _scalar_result(value):
    res = MagicMock()
    res.scalar_one_or_none.return_value = value
    return res


def _risk_card_db(ent, objs, detail_obj=None, events_obj=None):
    """按 SQL 文本特征分发查询（同 test_risk_notice_card_api.py 模式）。"""
    db = AsyncMock()
    db.add = MagicMock()

    def fake_execute(stmt):
        text = str(stmt)
        if "FROM enterprises" in text:
            return _scalar_result(ent)
        if "risk_notice_cards" in text:
            res = MagicMock()
            res.scalars.return_value.first.return_value = None
            return res
        if "FROM risk_objects" in text:
            if "enterprise_id =" in text:
                if "ORDER BY" in text:
                    return _rows_result(objs)
                return _scalar_result(detail_obj)
            return _scalar_result(events_obj)
        return _rows_result([])

    db.execute.side_effect = fake_execute
    return db


def _export_db(ent, obj, valid_ids):
    """导出用 mock：按绑定参数区分「存在/不存在」的 object_id。"""
    db = AsyncMock()
    db.add = MagicMock()

    def fake_execute(stmt):
        text = str(stmt)
        if "FROM enterprises" in text:
            return _scalar_result(ent)
        if "risk_notice_cards" in text:
            res = MagicMock()
            res.scalars.return_value.first.return_value = None
            return res
        if "FROM risk_objects" in text:
            params = stmt.compile().params
            id_params = [v for k, v in params.items() if k.startswith("id_")]
            oid = id_params[0] if id_params else None
            if "enterprise_id =" in text:
                if "ORDER BY" in text:
                    return _rows_result([obj])
                return _scalar_result(obj if oid in valid_ids else None)
            return _scalar_result(obj)
        return _rows_result([])

    db.execute.side_effect = fake_execute
    return db


@pytest.fixture()
def client(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "EXPORT_DIR", str(tmp_path))
    app = FastAPI()
    app.include_router(risk_notice_card.router)

    current_user = User(id="u1", email="a@b.c", name="A", role="admin")

    def _override_user():
        return current_user

    app.dependency_overrides[get_current_user] = _override_user
    app.dependency_overrides[get_db] = lambda: _risk_card_db(None, [])
    with TestClient(app) as test_client:
        yield test_client


def test_export_returns_file_key_and_writes_docx(client):
    """导出端点集成：真实 SVG→PNG 渲染，docx 落盘且含图片与标题。"""
    ent = _enterprise()
    obj = _risk_object()
    obj.events.append(_fire_event())
    client.app.dependency_overrides[get_db] = lambda: _risk_card_db(
        ent, [obj], detail_obj=obj, events_obj=obj
    )

    resp = client.post(
        "/enterprises/e1/risk-notice-cards/export", json={"object_ids": ["o1"]}
    )
    assert resp.status_code == 200
    data = resp.json()["data"]
    assert data["file_key"].startswith("risk-notice-e1-")
    assert data["file_key"].endswith(".docx")
    assert data["warnings"] == []

    out = Path(settings.EXPORT_DIR) / data["file_key"]
    assert out.is_file()
    doc = Document(str(out))
    texts = _all_paragraph_texts(doc)
    assert any("配电室安全风险告知卡" in t for t in texts)
    assert _drawings(doc) >= 2  # 二维码 + 安全标志
    # 真实渲染尺寸校验：内嵌 PNG 宽高均 > 1，避免 1x1 占位图掩盖渲染失败
    png_sizes = _embedded_png_sizes(doc)
    assert len(png_sizes) >= 2
    assert all(w > 1 and h > 1 for w, h in png_sizes)
    # 二维码实际内容为完整公开页 URL（导出端点由 request.base_url 推导）
    assert "http://testserver/r/tok1" in _embedded_qr_contents(doc)


def test_export_skips_missing_object_with_warnings(client, monkeypatch):
    async def _fake_svg_to_png(svg_name):
        return PLACEHOLDER_PNG

    monkeypatch.setattr(risk_notice_card, "svg_to_png", _fake_svg_to_png)
    ent = _enterprise()
    obj = _risk_object()
    obj.events.append(_fire_event())
    client.app.dependency_overrides[get_db] = lambda: _export_db(
        ent, obj, valid_ids={"o1"}
    )

    resp = client.post(
        "/enterprises/e1/risk-notice-cards/export",
        json={"object_ids": ["missing", "o1"]},
    )
    assert resp.status_code == 200
    data = resp.json()["data"]
    assert data["file_key"]
    assert any("missing" in w for w in data["warnings"])
    assert (Path(settings.EXPORT_DIR) / data["file_key"]).is_file()


def test_export_all_invalid_returns_400(client):
    ent = _enterprise()
    client.app.dependency_overrides[get_db] = lambda: _risk_card_db(
        ent, [], detail_obj=None
    )

    resp = client.post(
        "/enterprises/e1/risk-notice-cards/export",
        json={"object_ids": ["nope"]},
    )
    assert resp.status_code == 400
    assert "没有可导出的卡片" in resp.json()["detail"]
